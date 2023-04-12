from docx import Document
from docx import Document
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


document = Document()

sections = document.sections
    
sec_pr = document.sections[0]._sectPr # get the section properties el
# create new borders el
pg_borders = OxmlElement('w:pgBorders')
# specifies how the relative positioning of the borders should be calculated
pg_borders.set(qn('w:offsetFrom'), 'page')
for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
    border_el = OxmlElement(f'w:{border_name}')
    border_el.set(qn('w:val'), 'single') # a single line
    border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
    border_el.set(qn('w:space'), '24')
    border_el.set(qn('w:color'), 'auto')
    pg_borders.append(border_el) # register single border to border el
sec_pr.append(pg_borders) # apply border changes to section



paragraph=document.add_heading("M.H. Saboo Siddik College of Engineering",level=1)
font = paragraph.style.font
font.size=Pt(17)
font.underline = True
paragraph.alignment=1


paragraph=document.add_paragraph('"Deaprtment of Information Technology"', style='Intense Quote')
paragraph.alignment=1

document.add_paragraph(
    'Semester: III'
)
document.add_paragraph(
   'Room:205/208'
)

timetable=[{'time': '9.00am-10.00am  ', 'M': 'BC g VB / CC g ZM', 'T': '--', 'W': 'CCS L ZM', 'Th': '--', 'F': '--'}, {'time': '10.00am-11.00am  ', 'M': 'BDLT L VB', 'T': 'BDLT L VB', 'W': 'EM L AW', 'Th': '--', 'F': 'BC g VB / CC g ZM'}, {'time': '11.00am-12.00am', 'M': 'BDA L AS', 'T': 'BC g VB / CC g ZM', 
'W': '--', 'Th': 'BDA L AS', 'F': 'BDA L AS'}, {'time': '12.00am-1.00am  ', 'M': '--', 'T': '--', 'W': '--', 'Th': '--', 'F': '--'}, {'time': '2.00am-3.00am  ', 'M': '--', 'T': 'EM L AW', 'W': '--', 'Th': 'CCS L ZM', 'F': 'CCS L ZM'}, {'time': '3.00am-4.00am  ', 'M': 'EM L AW', 'T': 'CCS L ZM', 'W': 'BDLT L VB', 'Th': 'EM L AW', 'F': '--'}]


# main time table
table = document.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Day/Time'
hdr_cells[1].text = 'Monday'
hdr_cells[2].text = 'Tuesday'
hdr_cells[3].text = 'Wednesday'
hdr_cells[4].text = 'Thursday'
hdr_cells[5].text = 'Friday'

for records in timetable:
    row_cells = table.add_row().cells
    row_cells[0].text = records['time']
    row_cells[1].text = records['M']
    row_cells[2].text = records['T']
    row_cells[3].text = records['W']
    row_cells[4].text = records['Th']
    row_cells[5].text = records['F']

document.add_page_break()

document.save('docx\ timetable with heading.docx')