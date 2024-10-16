import os
import uuid
from app.s3_utils import download_image_from_s3
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import current_app as app
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL


def generatePaper(questions, data):
    document = Document()

    # Formating all sections
    sections = document.sections
    for section in sections:
        section.footer_distance = Inches(0.15)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("HeadingStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = "Times New Roman"

    # function for Fontstyle
    def fontstyle(
        paragraph,
        font_name="Times New Roman",
        font_size=14,
        font_bold=True,
        font_italic=False,
        font_underline=False,
    ):
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline

    # Function for Write Paragrapgh in word
    def writedocx(
        content,
        font_name="Times New Roman",
        font_size=12,
        font_bold=False,
        font_italic=False,
        font_underline=False,
        color=RGBColor(0, 0, 0),
        before_spacing=2,
        after_spacing=4,
        line_spacing=1.5,
        keep_together=True,
        keep_with_next=False,
        page_break_before=False,
        widow_control=False,
        align="left",
        style="List Number",
    ):
        paragraph = document.add_paragraph(str(content))
        # paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
        paragraph.style = document.styles[style]
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline
        font.color.rgb = color
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control
        if align.lower() == "left":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align.lower() == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align.lower() == "right":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align.lower() == "justify":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def insert_images_in_grid(image_paths):
        # Create a table with two rows and two columns
        table = document.add_table(rows=2, cols=2)

        # Set the alignment of the cells to center
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Cm(4.0)
                cell.height = Cm(2)

        # Define the desired cell dimensions (in centimeters)
        cell_width = Cm(5)
        cell_height = Cm(1)

        # Set the dimensions of each cell
        for row in table.rows:
            for cell in row.cells:
                cell.width = cell_width
                cell.height = cell_height

        # Iterate through the cells and insert images
        for i, cell in enumerate(table._cells):
            label = f"({i + 1})"
            cell_paragraph = cell.add_paragraph(label)
            cell_paragraph.alignment = 0  # Left alignment
            run = cell_paragraph.runs[0]
            run.bold = False

            # image_stream = download_image_from_s3(image_paths[i])

            # Add the image to the cell
            cell_paragraph = cell.add_paragraph()
            run = cell_paragraph.add_run()
            print("................",image_paths[i])
            # run.add_picture(image_stream, width=Cm(4.0), height=Cm(4.0))
            run.add_picture(image_paths[i], width=Cm(4.0), height=Cm(4.0))

        # # Minimize the gap between rows by adjusting row heights
        # for row in table.rows:
        #     row.height = Cm(3.2)  # Adjust as needed

    # Function for Footer page Number
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        page_run = paragraph.add_run()
        t1 = create_element("w:t")
        create_attribute(t1, "xml:space", "preserve")
        t1.text = "Page "
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = create_element("w:fldChar")
        create_attribute(fldChar1, "w:fldCharType", "begin")

        instrText = create_element("w:instrText")
        create_attribute(instrText, "xml:space", "preserve")
        instrText.text = "PAGE"

        fldChar2 = create_element("w:fldChar")
        create_attribute(fldChar2, "w:fldCharType", "end")

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = create_element("w:t")
        create_attribute(t2, "xml:space", "preserve")
        t2.text = " of "
        of_run._r.append(t2)

        fldChar3 = create_element("w:fldChar")
        create_attribute(fldChar3, "w:fldCharType", "begin")

        instrText2 = create_element("w:instrText")
        create_attribute(instrText2, "xml:space", "preserve")
        instrText2.text = "NUMPAGES"

        fldChar4 = create_element("w:fldChar")
        create_attribute(fldChar4, "w:fldCharType", "end")

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

    # craeting first section for header
    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")

    paragraph = document.add_heading(data["institute_name"])
    paragraph.alignment = 1
    fontstyle(paragraph, font_size=28, font_underline=True)

    # header table
    htable = document.add_table(1, 3)

    htab_cells = htable.rows[0].cells

    ht0 = htab_cells[0].add_paragraph(f"Exam : {data['exam']}\nDate : {data['date']}")
    # fontstyle(ht0)
    ht0.alignment = 0

    if data.get("set"):
        ht1 = htab_cells[1].add_paragraph(f"{data['paper_name']}\nSET : {data['set']}")
    else:
        ht1 = htab_cells[1].add_paragraph(f"{data['paper_name']}")
    ht1.alignment = 1

    ht2 = htab_cells[2].add_paragraph(
        f"Marks : {data['marks']}\nTime : {data['time']} min"
    )
    ht2.alignment = 2

    line = document.add_paragraph(
        "_________________________________________________________________________________"
    )

    # creating a main section
    main_section = document.add_section(0)
    main_section.top_margin = Inches(0.3)
    main_section.is_linked_to_previous = False

    # creating two column in page
    sectPr = main_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    # paragraph=document.add_paragraph()
    # paragraph.add_run(f"                  ---Set : {data['set']}---", style = 'HeadingStyle').bold = True
    # paragraph.alignment=0

    # calling writedocx function with for loop
    for i in questions:
        question = i["question"]
        question = question.replace("\n", " ") + f"  [{i['question_code']}]"

        if i["image_url"]:
            writedocx(question)
            img_question = document.add_paragraph()
            r = img_question.add_run()
            # image_stream = download_image_from_s3(i["image_url"])
            # r.add_picture(image_stream, width=Cm(6))
            r.add_picture(i["image_url"], width=Cm(6))
            img_question.alignment = 1

        if i["options"][0]["option_text"] and i["image_url"]:
            a1 = i["options"][0]["option_text"]
            a2 = i["options"][1]["option_text"]
            a3 = i["options"][2]["option_text"]
            a4 = i["options"][3]["option_text"]
            if len(a1) >= 26 or len(a2) >= 26 or len(a3) >= 26 or len(a4) >= 26:
                sentence = (
                    "\n" + "(1) " + a1 + "\n(2) " + a2 + "\n(3) " + a3 + "\n(4) " + a4
                )
            elif (
                13 < len(a1) <= 25
                or 13 < len(a2) <= 25
                or 13 < len(a3) <= 25
                or 13 < len(a4) <= 25
            ):
                sentence = (
                    "\n" + "(1) " + a1 + "\t(2) " + a2 + "\n(3) " + a3 + "\t(4) " + a4
                )
            else:
                sentence = (
                    "\n" + "(1) " + a1 + "  (2) " + a2 + "  (3) " + a3 + "  (4) " + a4
                )
            writedocx(sentence + "\n", style="List Bullet")

            # # Create a custom paragraph style without numbering
            # paragraph = document.add_paragraph(sentence)
            # custom_style = document.styles.add_style('CustomStyle', 1)  # Style ID is 1
            # custom_style.paragraph_format.left_indent = Cm(0.5)  # Adjust as needed
            # custom_style.paragraph_format.alignment = 1  # Center alignment
            # custom_style.font.bold = True

        elif i["options"][0]["option_text"]:
            a1 = i["options"][0]["option_text"]
            a2 = i["options"][1]["option_text"]
            a3 = i["options"][2]["option_text"]
            a4 = i["options"][3]["option_text"]
            if len(a1) >= 26 or len(a2) >= 26 or len(a3) >= 26 or len(a4) >= 26:
                sentence = (
                    question
                    + "\n"
                    + "(1) "
                    + a1
                    + "\n(2) "
                    + a2
                    + "\n(3) "
                    + a3
                    + "\n(4) "
                    + a4
                )
            elif (
                13 < len(a1) <= 25
                or 13 < len(a2) <= 25
                or 13 < len(a3) <= 25
                or 13 < len(a4) <= 25
            ):
                sentence = (
                    question
                    + "\n"
                    + "(1) "
                    + a1
                    + "\t(2) "
                    + a2
                    + "\n(3) "
                    + a3
                    + "\t(4) "
                    + a4
                )
            else:
                sentence = (
                    question
                    + "\n"
                    + "(1) "
                    + a1
                    + "  (2) "
                    + a2
                    + "  (3) "
                    + a3
                    + "  (4) "
                    + a4
                )
            writedocx(sentence + "\n")

        elif i["options"][0]["image_url_option"] and i["image_url"]:
            a1 = i["options"][0]["image_url_option"]
            a2 = i["options"][1]["image_url_option"]
            a3 = i["options"][2]["image_url_option"]
            a4 = i["options"][3]["image_url_option"]
            insert_images_in_grid([a1, a2, a3, a4])

        elif i["options"][0]["image_url_option"]:
            a1 = i["options"][0]["image_url_option"]
            a2 = i["options"][1]["image_url_option"]
            a3 = i["options"][2]["image_url_option"]
            a4 = i["options"][3]["image_url_option"]
            writedocx(question)
            insert_images_in_grid([a1, a2, a3, a4])

    # Calling for Page number in foooter
    footer = document.sections[0]
    add_page_number(footer.footer.paragraphs[0])

    # add new section
    new_section = document.add_section(0)
    # new_section.start_type=WD_SECTION.NEW_PAGE
    sectPr = new_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")

    endpara = document.add_paragraph("\n---------------Best of Luck---------------")
    fontstyle(endpara)
    endpara.alignment = 1

    # filename = "QuestionPaper.docx"

    document.save('docx/QuestionPaperwithImage.docx')

    return 


def answerKey(questions, data):
    document = Document()

    # Formating all sections
    sections = document.sections
    for section in sections:
        section.footer_distance = Inches(0.15)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)

    # function for Fontstyle
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style("HeadingStyle", WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = "Times New Roman"

    # Function for Write Paragrapgh in word
    def writedocx(
        content,
        font_name="Times New Roman",
        font_size=12,
        font_bold=False,
        font_italic=False,
        font_underline=False,
        color=RGBColor(0, 0, 0),
        before_spacing=2,
        after_spacing=4,
        line_spacing=1.5,
        keep_together=True,
        keep_with_next=False,
        page_break_before=False,
        widow_control=False,
        align="left",
        style="Normal",
    ):
        paragraph = document.add_paragraph(str(content))
        # paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
        paragraph.style = document.styles[style]
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline
        font.color.rgb = color
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control
        if align.lower() == "left":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align.lower() == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align.lower() == "right":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align.lower() == "justify":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Function for Footer page Number
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        page_run = paragraph.add_run()
        t1 = create_element("w:t")
        create_attribute(t1, "xml:space", "preserve")
        t1.text = "Page "
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = create_element("w:fldChar")
        create_attribute(fldChar1, "w:fldCharType", "begin")

        instrText = create_element("w:instrText")
        create_attribute(instrText, "xml:space", "preserve")
        instrText.text = "PAGE"

        fldChar2 = create_element("w:fldChar")
        create_attribute(fldChar2, "w:fldCharType", "end")

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = create_element("w:t")
        create_attribute(t2, "xml:space", "preserve")
        t2.text = " of "
        of_run._r.append(t2)

        fldChar3 = create_element("w:fldChar")
        create_attribute(fldChar3, "w:fldCharType", "begin")

        instrText2 = create_element("w:instrText")
        create_attribute(instrText2, "xml:space", "preserve")
        instrText2.text = "NUMPAGES"

        fldChar4 = create_element("w:fldChar")
        create_attribute(fldChar4, "w:fldCharType", "end")

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

    # craeting first section for header
    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")

    paragraph = document.add_heading(data["institute_name"])
    font = paragraph.style.font
    font.size = Pt(25)
    font.underline = True

    paragraph.alignment = 1
    # fontstyle(paragraph, font_size=28,font_underline=True)

    # header table
    htable = document.add_table(1, 3)

    htab_cells = htable.rows[0].cells

    ht0 = htab_cells[0].add_paragraph()
    ht0.add_run(
        f"Exam : {data['exam']}\nDate : {data['date']}", style="HeadingStyle"
    ).bold = True
    # fontstyle(ht0,font_bold=True)
    ht0.alignment = 0

    ht1 = htab_cells[1].add_paragraph()
    if data.get("set"):
        ht1.add_run(
            f"{data['paper_name']}\n SET : {data['set']}", style="HeadingStyle"
        ).bold = True
    else:
        ht1.add_run(f"{data['paper_name']}", style="HeadingStyle").bold = True

    ht1.alignment = 1

    ht2 = htab_cells[2].add_paragraph()
    ht2.add_run(
        f"Marks : {data['marks']}\nTime : {data['time']} min", style="HeadingStyle"
    ).bold = True

    ht2.alignment = 2

    document.add_paragraph(
        "________________________________________________________________________________________________________________________________________"
    ).bold = True
    paragraph = document.add_paragraph()
    paragraph.add_run(f"-----Answer-----", style="HeadingStyle").bold = True
    paragraph.alignment = 1

    # creating a main section
    main_section = document.add_section(0)
    main_section.top_margin = Inches(0.3)
    main_section.is_linked_to_previous = False

    # creating two column in page
    sectPr = main_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "6")

    # calling writedocx function with for loop
    flag = 0
    for i in questions:
        answer = i["answer"]
        exp = i["explaination"]
        if exp:
            flag = 1
        writedocx(
            "(" + answer + ")\n", font_bold=True, style="List Number", font_size=13
        )

    # Calling for Page number in foooter
    footer = document.sections[0]
    add_page_number(footer.footer.paragraphs[0])

    if flag:
        document.add_page_break()

        line_section = document.add_section(0)
        sectPr = line_section._sectPr
        cols = sectPr.xpath("./w:cols")[0]

        cols.set(qn("w:num"), "1")

        paragraph = document.add_paragraph()
        paragraph.add_run("Hint & Solution", style="HeadingStyle").bold = True
        paragraph.alignment = 1

        # add new section
        new_section = document.add_section(0)
        new_section.is_linked_to_previous = False
        sectPr = new_section._sectPr
        cols = sectPr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), "2")

        lineNo = 1
        for i in questions:
            answer = i["answer"]
            exp = i["explaination"]
            if exp:
                p = document.add_paragraph()
                p.add_run(f"{lineNo}.").bold = True
                p.add_run(f" Answer ({answer})\n")
                p.add_run("Sol. ").bold = True
                # writedocx(f"{lineNo}. Answer ({answer})\nSol.")
                p.add_run(exp)
                # fontstyle(p,font_bold=False,font_size=12)

                # writedocx(exp,font_size=12.5)
            else:
                pass
            lineNo += 1

    new_section = document.add_section(0)
    new_section.is_linked_to_previous = False
    sectPr = new_section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")
    endpara = document.add_paragraph()
    endpara.add_run(
        "---------------End---------------", style="HeadingStyle"
    ).bold = True
    # fontstyle(endpara)
    endpara.alignment = 1

    filename = "AnswerKey.docx"

    # path = os.path.join(
    #     app.config["UPLOAD_FOLDER"], "{}.{}".format(str(uuid.uuid4()), filename)
    # )

    document.save('docx/AnswerKey.docx')
    return 
