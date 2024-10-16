# import matplotlib.pyplot as plt
# from docx import Document
# from docx.shared import Inches

# # Step 1: Generate the LaTeX equation as an image
# def generate_latex_image(latex_code, image_filename):
#     plt.figure(figsize=(4, 1))  # Adjusted size for complex formulas
#     plt.text(0.5, 0.5, f"${latex_code}$", fontsize=20, ha='center', va='center')
#     plt.axis('off')
#     plt.savefig(image_filename, bbox_inches='tight', pad_inches=0.1)
#     plt.close()

# # LaTeX equations to render
# latex_equation1 = r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}"  # Quadratic formula
# latex_equation2 = r"\oint_V f(s) \, ds"  # Contour integral
# latex_equation3 = r"\text{C}_6\text{H}_{12} + 6\text{O}_2 \rightarrow 6\text{CO}_2 + 6\text{H}_2\text{O}"  # Combustion of glucose
# latex_equation5 = r"\int_{a}^{b} x^2 \,dx"  # Definite integral

# # Image filenames
# image_filename1 = "latex_equation1.png"
# image_filename2 = "latex_equation2.png"
# image_filename3 = "latex_equation3.png"
# image_filename5 = "latex_equation5.png"

# # Generate images for the equations
# generate_latex_image(latex_equation1, image_filename1)
# generate_latex_image(latex_equation2, image_filename2)
# generate_latex_image(latex_equation3, image_filename3)
# generate_latex_image(latex_equation5, image_filename5)


# # Step 2: Create a Word document and insert the images
# document = Document()
# document.add_heading('Document Title', 0)

# # Add the first LaTeX image to the document
# document.add_paragraph('Here is the first LaTeX equation (Quadratic Formula):')
# document.add_picture(image_filename1, width=Inches(3))  # Adjust the width as needed

# # Add the second LaTeX image to the document
# document.add_paragraph('Here is the second LaTeX equation (Contour Integral):')
# document.add_picture(image_filename2, width=Inches(3))  # Adjust the width as needed

# # Add the third LaTeX image (chemistry formula) to the document
# document.add_paragraph('Here is a complex chemistry formula (Combustion of Glucose):')
# document.add_picture(image_filename3, width=Inches(4))  # Adjust the width as needed

# document.add_paragraph('Here is a complex chemistry formula (Combustion of Glucose):')
# document.add_picture(image_filename5, width=Inches(4))  # Adjust the width as needed

# # Save the document
# document.save('demo_with_latex.docx')

# print("Word document created successfully with the LaTeX images.")


import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# Step 1: Generate the LaTeX equation as an image
def generate_latex_image(latex_code, image_filename):
    plt.figure(figsize=(2, 1))
    plt.text(0.5, 0.5, f"${latex_code}$", fontsize=20, ha='center', va='center')
    plt.axis('off')
    plt.savefig(image_filename, bbox_inches='tight', pad_inches=0.1)
    plt.close()

# LaTeX equation to render
latex_equation = r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}"
latex_equation = r"f(x) = a_0 + \sum_{n=1}^{\infty} \left( a_n \cos\left( \frac{n\pi x}{L} \right) + b_n \sin\left( \frac{n\pi x}{L} \right) \right)"
image_filename = "latex_equation.png"

generate_latex_image(latex_equation, image_filename)

# Step 2: Create a Word document and insert the image
document = Document()
document.add_heading('Document Title', 0)

# Add the LaTeX image to the document
document.add_paragraph('Here is the LaTeX equation:')
document.add_picture(image_filename, width=Inches(3))  # Adjust the width as needed

# Save the document
document.save('demo_with_latex.docx')

print("Word document created successfully with the LaTeX image.")
