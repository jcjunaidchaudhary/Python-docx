from docx import Document
from docx.shared import Inches

# ------- initial code -------

document = Document()

p = document.add_paragraph()
r = p.add_run()
r.add_text('Good Morning every body,This is my ')
picPath = 'static/my dp.png'
r.add_picture(picPath)
r.add_text(' do you like it?')

document.save('docx\writeWithPicture.docx')

# ------- improved code -------

document = Document()

p = document.add_paragraph('Picture bullet section', 'List Bullet')
p = p.insert_paragraph_before('')
r = p.add_run()
r.add_picture(picPath)
p = p.insert_paragraph_before('My picture title', 'Heading 1')

document.save('docx\writeWithPicture2.docx')