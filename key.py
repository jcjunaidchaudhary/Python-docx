from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def paperKey(obj):
    document = Document()

# Formating all sections
    sections = document.sections
    for section in sections:
        section.footer_distance=Inches(0.15)
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)


    #function for Fontstyle
    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('HeadingStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font   
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    #Function for Write Paragrapgh in word
    def writedocx(content, font_name = 'Times New Roman', font_size = 12, font_bold = False, font_italic = False, font_underline = False, color = RGBColor(0, 0, 0),
                before_spacing = 2, after_spacing = 4, line_spacing = 1.5, keep_together = True, keep_with_next = False, page_break_before = False,
                widow_control = False, align = 'left', style = 'Normal'):
        paragraph = document.add_paragraph(str(content))
        #paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
        paragraph.style = document.styles[style]
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline
        font.color.rgb = color
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control
        if align.lower() == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align.lower() == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align.lower() == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align.lower() == 'justify':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    #Function for Footer page Number
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        page_run = paragraph.add_run()
        t1 = create_element('w:t')
        create_attribute(t1, 'xml:space', 'preserve')
        t1.text = 'Page '
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = create_element('w:t')
        create_attribute(t2, 'xml:space', 'preserve')
        t2.text = ' of '
        of_run._r.append(t2)

        fldChar3 = create_element('w:fldChar')
        create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = create_element('w:instrText')
        create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = create_element('w:fldChar')
        create_attribute(fldChar4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)


    #craeting first section for header
    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

    paragraph=document.add_heading("Radiance Academy")
    font = paragraph.style.font
    font.size=Pt(25)
    font.underline = True

    paragraph.alignment=1
    # fontstyle(paragraph, font_size=28,font_underline=True)

    #header table
    htable=document.add_table(1,3)

    htab_cells=htable.rows[0].cells

    ht0=htab_cells[0].add_paragraph()
    ht0.add_run("Exam : Neet\nDate : 12/2/22",style="HeadingStyle").bold=True
    # fontstyle(ht0,font_bold=True)
    ht0.alignment = 0
    
    ht1=htab_cells[1].add_paragraph()
    ht1.add_run("Mock Test\nNEET",style="HeadingStyle").bold=True

    ht1.alignment = 1

    ht2=htab_cells[2].add_paragraph()
    ht2.add_run("Marks : 720\nTime : 3 hour",style="HeadingStyle").bold=True

    ht2.alignment = 2

    document.add_paragraph("________________________________________________________________________________________________________________________________________")

    #creating a main section
    main_section = document.add_section(0)
    main_section.top_margin=Inches(0.3)
    main_section.is_linked_to_previous=False

    #creating two column in page
    sectPr = main_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '7')
    

    #calling writedocx function with for loop
    flag=0
    for i in obj:
        answer=i["answer"]
        exp=i["explaination"]
        if exp:
            flag=1
        writedocx("("+answer+")\n", font_bold=True,style='List Number',font_size=13)
    
    # Calling for Page number in foooter
    footer=document.sections[0]
    add_page_number(footer.footer.paragraphs[0])
    
    
    
    if flag:
        document.add_page_break()

        line_section = document.add_section(0)
        sectPr = line_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        
        cols.set(qn('w:num'), '1')
    
        paragraph=document.add_paragraph()
        paragraph.add_run("---------------------Explanation------------------", style = 'HeadingStyle').bold = True
        paragraph.alignment=1
    
    

        #add new section
        new_section = document.add_section(0)
        new_section.is_linked_to_previous=False
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
    
        lineNo=1
        for i in obj:
            answer=i["answer"]
            exp=i["explaination"]
            if exp:
                p = document.add_paragraph()
                p.add_run(f"{lineNo}.").bold = True
                p.add_run(f' Answer ({answer})\n')
                p.add_run('Sol. ').bold = True
                # writedocx(f"{lineNo}. Answer ({answer})\nSol.")
                p.add_run(exp)
                # fontstyle(p,font_bold=False,font_size=12)

                # writedocx(exp,font_size=12.5)
            else:
                pass
            lineNo+=1


    new_section = document.add_section(0)
    new_section.is_linked_to_previous=False
    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')
    endpara=document.add_paragraph()
    endpara.add_run("---------------End---------------",style="HeadingStyle").bold=True
    # fontstyle(endpara)
    endpara.alignment=1

    document.save('key.docx')



#Objects
lst=[{  "question":"The acceleration ‘a’ in m/s²  of a particle is given by a = 3t²  + 2t + 2 where t is the time. If the particle starts out with a velocity  u = 2m /s  at t = 0, then the velocity at the end of  2 second is.",
            "Option1":"12 m/s",
            "option2":"18 m/s",
            "optiin3":"27 m/s",
            "option4":"36 m/s",
            "answer":"2"  ,
            "explaination" :"" 
        },
        {  "question":"The work done in an adiabatic change in a gas depends only on.",
            "Option1":"Change is pressure",
            "option2":"Change in volume",
            "optiin3":"change in temprature",
            "option4":"None of these" ,
            "answer":"3"   ,
            "explaination" :"floats with ¼ of its volume above the water level."
        },
        {  "question":"In the case of constants  and α of β a transistor.",
            "Option1":"1.2",
            "option2":"441",
            "optiin3":"444",
            "option4":"433"  ,
            "answer":"1"  ,
            "explaination" :""
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique" ,  
            "answer":"5" ,
            "explaination" :""
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique" ,  
            "answer":"5" ,
            "explaination" :""
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"4" ,
            "explaination" :"floats with ¼ of its volume above the water level."   
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",   
            "answer":"1" ,
            "explaination" :""
        },
        {  "question":"The work done in an adiabatic change in a gas depends only on.",
            "Option1":"Change is pressure",
            "option2":"Change in volume",
            "optiin3":"change in temprature",
            "option4":"None of these",
            "answer":"1" ,
            "explaination" :""   
        },
        {  "question":"In the case of constants  and α of β a transistor.",
            "Option1":"1.2",
            "option2":"441",
            "optiin3":"444",
            "option4":"433" ,   
            "answer":"2",
            "explaination" :""
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :"this is beacuse aszhraf is malik with out having ek koadi"
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :""
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"my name is chaudhary not khan" 
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :"this is beacuse aszhraf is malik with out having ek koadi"
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :"sssd"
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"my name is chaudhary not khan" 
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :""
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :"sssd"
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"my name is chaudhary not khan" 
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :"this is beacuse aszhraf is malik with out having ek koadi"
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :"sssd"
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"" 
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :"this is beacuse aszhraf is malik with out having ek koadi"
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :""
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"my name is chaudhary not khan" 
        },
        {  "question":"What is your name ",
            "Option1":"Ashraf",
            "option2":"Junaid",
            "optiin3":"Suhail",
            "option4":"Sadique",
            "answer":"3"    ,
            "explaination" :"this is beacuse aszhraf is malik with out having ek koadi"
        },
        {  "question":"Mumbai coding club is initialzed by.",
            "Option1":"Gani bhai",
            "option2":"Majnju Bhai",
            "optiin3":"Uday bhai",
            "option4":"Babu rao aapte",
            "answer":"1"    ,
            "explaination" :"sssd"
        },
        {  "question":"B floats with ¼ of its volume above the water level.",
            "Option1":"hii",
            "option2":"2:3",
            "optiin3":"4:5",
            "option4":"43",  
            "answer":"1" ,
            "explaination" :"my name is chaudhary not khan" 
        }
        
]

# lst=[{"question":"The acceleration ‘a’ in m/s²  of a particle is given by a = 3t²  + 2t + 2 where t is the time. If the particle starts out with a velocity  u = 2m /s  at t = 0, then the velocity at the end of  2 second is.",
#             "Option1":"12 m/s",
#             "option2":"18 m/s",
#             "optiin3":"27 m/s",
#             "option4":"36 m/s",
#             "answer":"2"  ,
#             "explaination" :"" 
#         },
#         {  "question":"The work done in an adiabatic change in a gas depends only on.",
#             "Option1":"Change is pressure",
#             "option2":"Change in volume",
#             "optiin3":"change in temprature",
#             "option4":"None of these" ,
#             "answer":"3"   ,
#             "explaination" :""
#         },
#         {  "question":"In the case of constants  and α of β a transistor.",
#             "Option1":"1.2",
#             "option2":"441",
#             "optiin3":"444",
#             "option4":"433"  ,
#             "answer":"1"  ,
#             "explaination" :""
#         },
#         {  "question":"What is your name ",
#             "Option1":"Ashraf",
#             "option2":"Junaid",
#             "optiin3":"Suhail",
#             "option4":"Sadique" ,  
#             "answer":"1" ,
#             "explaination" :""
#         }]
paperKey(lst)