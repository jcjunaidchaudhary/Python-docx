from docx import Document
from docx.enum.style import WD_STYLE_TYPE 
from docx.shared import Pt, RGBColor, Inches
# from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.enum.section import WD_SECTION

document = Document()

# create section
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

section = document.sections[0]
# section.start_type = WD_SECTION.NEW_PAGE


#function for Fontstyle
def fontstyle(paragraph, font_name = 'Times New Roman', font_size = 12, font_bold = True, font_italic = False, font_underline = False):
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline

#Function for Write Paragrapgh in word
def writedocx(content, font_name = 'Times New Roman', font_size = 12, font_bold = False, font_italic = False, font_underline = False, color = RGBColor(0, 0, 0),
              before_spacing = 2, after_spacing = 4, line_spacing = 1.5, keep_together = True, keep_with_next = False, page_break_before = False,
              widow_control = False, align = 'left', style = ''):
    paragraph = document.add_paragraph(str(content))
    #paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    paragraph.style = document.styles['List Number']
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    paragraph.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control
    if align.lower() == 'left':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

#Function for Footer page Number
def create_element(name):
    return 

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    return
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


# Header
header = document.sections[0].header

paragraph = header.paragraphs[0]
paragraph.text = "Radiance Academy"
paragraph.alignment=1
fontstyle(paragraph, font_size=25,font_underline=True)

#header table
htable=header.add_table(1,3, Inches(8))

htab_cells=htable.rows[0].cells
ht0=htab_cells[0].add_paragraph("Exam : Neet\nDate : 12/2/22")
fontstyle(ht0)
ht0.alignment = 0

ht1=htab_cells[1].add_paragraph('Mock Test\nNEET')
ht1.alignment = 1

ht2=htab_cells[2].add_paragraph("Marks : 720\nTime : 3 hour")
ht2.alignment = 2


header.first_page_header =True
header.is_linked_to_previous = False


#creating two column in page
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')



#caliing with other file
# f=open('test.txt','r')
# for l in f:
#     writedocx(l)
 
#Objects
lst=[{  "question":"The acceleration ‘a’ in m/s²  of a particle is given by a = 3t²  + 2t + 2 where t is the time. If the particle starts out with a velocity  u = 2m /s  at t = 0, then the velocity at the end of  2 second is.",
        "Option1":"12 m/s",
        "option2":"18 m/s",
        "optiin3":"27 m/s",
        "option4":"36 m/s"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"The work done in an adiabatic change in a gas depends only on.",
        "Option1":"Change is pressure",
        "option2":"Change in volume",
        "optiin3":"change in temprature",
        "option4":"None of these"    
    },
    {  "question":"In the case of constants  and α of β a transistor.",
        "Option1":"1.2",
        "option2":"441",
        "optiin3":"444",
        "option4":"433"    
    },
    {  "question":"What is your name ",
        "Option1":"Ashraf",
        "option2":"Junaid",
        "optiin3":"Suhail",
        "option4":"Sadique"    
    },
    {  "question":"Mumbai coding club is initialzed by.",
        "Option1":"Gani bhai",
        "option2":"Majnju Bhai",
        "optiin3":"Uday bhai",
        "option4":"Babu rao aapte"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    },
    {  "question":" If r  represents the radius of the orbit of a satellite of mass m moving around a planet of mass M, the velocity of the satellite is given by",
        "Option1":"They are monochromatic ",
        "option2":"They are highly polarised    ",
        "optiin3":"They are coherent ",
        "option4":"They have high degree of parallelism"    
    },
    {  "question":"B floats with ¼ of its volume above the water level.",
        "Option1":"hii",
        "option2":"2:3",
        "optiin3":"4:5",
        "option4":"43"    
    }

]


#calling writedocx function with for loop
for i in lst:
    q=i["question"]
    a1=i["Option1"]
    a2=i["option2"]
    a3=i["optiin3"]
    a4=i["option4"]
    if len(a1)>=26 or len(a2)>=26 or len(a3)>=26 or len(a4)>=26:
        sentence=q+"\n"+"(1) "+a1+"\n(2) "+a2+"\n(3) "+a3+"\n(4) "+a4
    elif 13<len(a1)<=25 or 13<len(a2)<=25 or 13<len(a3)<=25 or 13<len(a4)<=25:
        sentence=q+"\n"+"(1) "+a1+"\t(2) "+a2+"\n(3) "+a3+"\t(4) "+a4
    else:
        sentence=q+"\n"+"(1) "+a1+"  (2) "+a2+"  (3) "+a3+"  (4) "+a4
    writedocx(sentence+"\n")

# Calling for Page number in foooter
add_page_number(document.sections[0].footer.paragraphs[0])

#add new section
new_section = document.add_section(0)
# new_section.start_type=WD_SECTION.NEW_PAGE
sectPr2 = new_section._sectPr
cols2 = sectPr2.xpath('./w:cols')[0]
cols2.set(qn('w:num'), '1')

endpara=document.add_paragraph("--------------Best of Luck---------------")
endpara.alignment=1

document.save('docx/QuestionPaperWithHeader.docx')

