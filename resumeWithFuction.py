import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX

document = Document()

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#function for Fontstyle
def fontstyle(paragraph, font_name = 'Times New Roman', font_size = 14, font_bold = False, font_italic = False, font_underline = False):
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline


def generateResume(personal,education,experience,projects,social,add_Info):
    
# create section
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)


    section = document.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    paragraph=document.add_heading(personal.name)
    paragraph.alignment=0
    font = paragraph.style.font
    font.size=Pt(20)
    font.underline = True

    par=document.add_paragraph('\n')
    par.add_run(f'Mobile No: {personal.mobile}').bold=True
    par.add_run(f'\nEmail: {personal.email}').bold=True
    par.add_run(f'\nLocation: {personal.city},{personal.state},{personal.country}').bold=True
    par.alignment=0

    pp=document.add_paragraph()
    r = pp.add_run()
    r.add_picture(personal.profile_picture, width=Inches(1.5))
    pp.alignment=2


    # fontstyle(paragraph, font_size=28,font_underline=True)
    main_section = document.add_section(0)
    main_section.top_margin=Inches(0.3)
    sectPr = main_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '0')

    document.add_heading('Objective')
    objective= document.add_paragraph(personal.objective)
    fontstyle(objective)

    add_section = document.add_section(0)
    add_section.top_margin=Inches(0.3)
    sectPr = add_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')

    document.add_heading('Work Experience')
    exp = document.add_paragraph()
    exp.add_run(f'{experience.designation}\n').bold = True
    exp.add_run(f'{experience.company_name},{experience.location}\n')
    exp.add_run(f'Working on following technologies in:\n{experience.workingOn}')
    exp.add_run(f'\n({experience.joining_date} - {experience.leaving_date})')

    document.add_heading('Academic Background')

    for ed in education:
        graduation = document.add_paragraph()
        graduation.add_run(f'{ed.education}').bold = True
        graduation.add_run(f'({ed.passing_year})').bold = True
        graduation.add_run(f'\n{ed.institute} | Score: {ed.percentage}')

    document.add_heading('Projects')

    for pr in projects:
        project=document.add_paragraph()
        project.add_run(f'{pr.project_name}\n').bold=True
        project.add_run(f'{pr.project_desc}')


    # skills="python, Java, Django, Flask, GitHub, Angular, HTML, CSS"
    document.add_heading('Technical Skills')
    skill=document.add_paragraph('')
    skill.add_run(add_Info.skills)

    document.add_heading('Personal Deatils') 
    persnal=document.add_paragraph(f'Date of birth : {personal.dob}\n')
    persnal.add_run(f"Marital Status : {add_Info.marital_Status}\n")
    persnal.add_run(f"Gender : Male\n")
    persnal.add_run(f"Language :{add_Info.language}")
    

    document.add_heading('Links') 

    for s in social:
        link=document.add_paragraph()
        link.add_run(f'{s.social_profile} ').bold=True
        add_hyperlink(link, f'{s.url}', f"{s.url}")


    website=document.add_paragraph()
    website.add_run("Website : ")
    add_hyperlink(website,f'{add_Info.website}', f"{add_Info.website}")


    document.save('docx\ResumeWithFunction.docx')



