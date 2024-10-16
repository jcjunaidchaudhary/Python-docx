from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document

def insert_images_in_single_column(document, image_paths):
    # Create a table with 4 rows and 1 column
    table = document.add_table(rows=4, cols=1)

    # Set the column width and vertical alignment for each cell
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5)  # Set consistent width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # Align content to the top of the cell

            # Clear any default cell padding by setting paragraph formatting
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    # Insert images in each row
    for i, cell in enumerate(table._cells):
        if i < len(image_paths):
            label = f"Image {i + 1}"  # Label for the image
            # Add label to the cell
            label_paragraph = cell.add_paragraph(label)
            label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align for label
            label_paragraph.paragraph_format.space_after = Pt(0)  # Remove extra space after label
            label_paragraph.paragraph_format.space_before = Pt(0)  # Remove extra space before label

            # Add the image below the label
            try:
                image_paragraph = cell.add_paragraph()
                run = image_paragraph.add_run()
                run.add_picture(image_paths[i], width=Cm(4.0))  # Set the width of the image
                image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the image

                # Remove extra spacing from the image paragraph
                image_paragraph.paragraph_format.space_after = Pt(0)
                image_paragraph.paragraph_format.space_before = Pt(0)
            except Exception as e:
                print(f"Error inserting image at index {i}: {e}")
        else:
            print(f"No image found for index {i}")

# Example usage
document = Document()
image_paths = ['latex_equation2.png', 'latex_equation1.png', 'latex_equation3.png', 'latex_equation4.png']  # List of image paths
insert_images_in_single_column(document, image_paths)
document.save('insert_images_in_single_column.docx')
