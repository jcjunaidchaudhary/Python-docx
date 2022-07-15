from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import true
document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

header = document.sections[0].header

def fontstyle(paragraph, font_name = 'Times New Roman', font_size = 12, font_bold = True, font_italic = False, font_underline = False):
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline


paragraph = header.paragraphs[0]
paragraph.text = "Radiance Academy"
paragraph.alignment=1
fontstyle(paragraph, font_size=25,font_underline=True)

htable=header.add_table(1,3, Inches(8))

htab_cells=htable.rows[0].cells
ht0=htab_cells[0].add_paragraph("Exam : Neet\nDate : 12/2/22")
fontstyle(ht0)
ht0.alignment = 0

ht1=htab_cells[1].add_paragraph('Mock Test\nNEET')
ht1.alignment = 1

ht2=htab_cells[2].add_paragraph("Marks : 720\nTime : 3 hour")
ht2.alignment = 2
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

insertHR(ht2)
insertHR(ht0)
insertHR(ht1)



document.save('header.docx')