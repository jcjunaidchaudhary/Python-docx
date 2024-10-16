import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches,Cm
from matplotlib import rc

# Function to generate LaTeX images
# def generate_latex_image(latex_code, image_filename):
#     plt.figure(figsize=(6, 1))
#     plt.text(0.5, 0.5, f"${latex_code}$", fontsize=18, ha='center', va='center')
#     plt.axis('off')
#     plt.savefig(image_filename, bbox_inches='tight', pad_inches=0)
#     plt.close()

def generate_latex_image(latex_code, image_filename):
    plt.figure(figsize=(2, 1))

    if latex_code.startswith('$'):
        plt.text(0.5, 0.5, latex_code, fontsize=20, ha='center', va='center',fontweight='normal', style='normal')
    else:
        plt.text(0.5, 0.5, f"${latex_code}$", fontsize=20, ha='center', va='center',fontweight='normal', style='normal')
    plt.axis('off')
    plt.savefig(image_filename, bbox_inches='tight', pad_inches=0.1)
    plt.close()



# Create a new document
document = Document()
document.add_heading('Question Paper', 0)

# Example question with LaTeX expressions
question = {
    'text': 'The acceleration ‘a’ in m/s² of a particle is given by',
    'latex': r"(x+a)^n = \sum_{k=0}^{n} \binom{n}{k} x^k a^{n-k}",
    'latex': r"f(x)=a_0+\sum_{n=1}^∞ (a_n  cos⁡ nπx/L +b_n  sin⁡ nπx/L )",
    'latex': r"sin⁡α±sin⁡β=2 sin⁡ 1/2 (α±β)  cos⁡ 1/2 (α∓β)",
    'latex': r"\sin α\pm \sin β=2 \sin \frac{1}{2} \left(α\pm β\right) cos⁡\frac{1}{2} \left(α\mp β\right)",
    'latex': r"f\left(x\right)=a_0+∑_\left(n=1\right)^∞ \left(a_n  cos⁡nπx/L+b_n  \sin nπx/L \right)",
    'latex': r"\sum_{n=1}^{\infty} 2^{-n} = 1",
    'latex': r"\sin \alpha \pm \sin \beta = 2 \sin\left(\frac{1}{2}(\alpha \pm \beta)\right) \cos\left(\frac{1}{2}(\alpha \mp \beta)\right)",
    'latex': r"\int_{a}^{b} x^2 \,dx",
    'end_text': 'where t is the time. If the particle starts out with a velocity u = 2m/s at t = 0, then the velocity at the end of 2 seconds is.',
    'options': [
        r"v = u + at",
        r"v = \frac{1}{2} at^2 + u",
        r"v = at^2 + u",
        "None of the above"
    ]
}


# Add question text
paragraph = document.add_paragraph(f"Q1: {question['text']} ")

# Generate LaTeX image for the formula and add it to the paragraph
generate_latex_image(question['latex'], "latex_question1.png")
paragraph.add_run().add_picture("latex_question1.png", width=Cm(4.0))  # Adjust width

# Add remaining question text after the LaTeX formula
document.add_paragraph(question['end_text'])

# Add options
for opt_idx, option in enumerate(question['options'], start=1):
    generate_latex_image(option, f"option_latex_{opt_idx}.png")
    document.add_paragraph(f"Option {chr(64+opt_idx)}: ")
    document.add_picture(f"option_latex_{opt_idx}.png", width=Inches(2.5))  # Adjust width if needed

# Save the document
document.save('question_paper_with_complex_latex2.docx')

print("Question Paper with complex LaTeX generated successfully.")
