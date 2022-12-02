import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX

document = Document()

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#function for Fontstyle
def fontstyle(paragraph, font_name = 'Times New Roman', font_size = 14, font_bold = False, font_italic = False, font_underline = False):
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline

# create section
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)


section = document.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')

paragraph=document.add_heading("Junaid Ahmad Chaudhary")
paragraph.alignment=0
font = paragraph.style.font
font.size=Pt(20)
font.underline = True

par=document.add_paragraph('\n')
par.add_run('Mobile No: +91 987654321').bold=True
par.add_run('\nEmail: jc.junaid.chaudhary@gmail.com').bold=True
par.add_run('\nLocation: Mumbai,Maharashtra,India').bold=True
par.alignment=0

pp=document.add_paragraph()
r = pp.add_run()
r.add_picture('static/my dp.png', width=Inches(1.5))
pp.alignment=2


# fontstyle(paragraph, font_size=28,font_underline=True)
main_section = document.add_section(0)
main_section.top_margin=Inches(0.3)
sectPr = main_section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '0')

document.add_heading('Objective')
objective= document.add_paragraph('Seeking the position of a Python Developer to further hone my skills in Python, Flask,Django, SQL Database to enhance organizational effectiveness')
fontstyle(objective)

add_section = document.add_section(0)
add_section.top_margin=Inches(0.3)
sectPr = add_section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')

document.add_heading('Work Experience')
exp = document.add_paragraph()
exp.add_run('Python Developer\n').bold = True
exp.add_run('Olx Information Technology,Jogeshwari Mumbai.\n')
exp.add_run('Working on following technologies in:\nPython3, Django, MySQL, REST API, HTML5,CCS3, JavaScript and GitHub.')
exp.add_run('\n(May 2022 - present)')

document.add_heading('Academic Background')

graduation = document.add_paragraph()
graduation.add_run('Bachelors of Engineering Information Technology').bold = True
graduation.add_run('(2023)').bold = True
graduation.add_run('\nM.H Saboo Siddik College of Engineering | CGPA: 8.29')

hsc = document.add_paragraph()
hsc.add_run('Ramniranjan Jhunjhunwala college | ').bold = True
hsc.add_run('HSC | (2019)').bold = True
hsc.add_run(' Score:64')

ssc = document.add_paragraph()
ssc.add_run('Yogiraj Shree Krishna Vidyalaya | ').bold = True
ssc.add_run('SSC | (2017)').bold = True
ssc.add_run(' Score:89')


document.add_heading('Projects')

projects=document.add_paragraph()
projects.add_run('Personalize Notebook\n').bold=True
# for word in lst:
#     projects.add_run(word)
projects.add_run('JSON based API, Based on Flask Framework, ORM tool used SQLALCHEMY Database Management Tool used MySQL | SQLit')


skills="python, Java, Django, Flask, GitHub, Angular, HTML, CSS"
document.add_heading('Technical Skills')
skill=document.add_paragraph('')
skill.add_run(skills)

document.add_heading('Personal Deatils') 
personal=document.add_paragraph('Date of birth : may 2020\n')
personal.add_run("Marital Status : Single\n")
personal.add_run("Gender : Male\n")
personal.add_run("Language :English, Hindi, Urdu")

document.add_heading('Links') 

link=document.add_paragraph()
link.add_run('LinkedIn: ').bold=True
add_hyperlink(link, 'https://www.linkedin.com/in/juned-ahmad-chaudhary-25498720a/', "https://www.linkedin.com/in/juned-ahmad-chaudhary-25498720a/")

git=document.add_paragraph()
git.add_run('GitHUb:  ').bold=True
add_hyperlink(git, 'https://github.com/jcjunaidchaudhary', "https://github.com/jcjunaidchaudhary")



document.save('docx\Resume.docx')