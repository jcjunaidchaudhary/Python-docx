from docx import Document
from docx.shared import Pt, Inches
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement, qn


document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

section = document.sections[0]

def set_no_borders(table):
    """ Set borders to zero width to effectively remove them. """
    for cell in table._element.iter_tcs():
        tcPr = cell.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            cell._insert_tcPr(tcPr)
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

        # List of the sides of the cell to apply no borders
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)

def set_table_center(table):
    """Set table alignment to center in its container."""
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        tbl_pr = tbl_pr[0]
    else:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    
    tbl_look = OxmlElement('w:tblLook')
    tbl_look.set(qn('w:val'), "04A0")
    tbl_pr.append(tbl_look)

    jc = OxmlElement('w:jc')  # Creates a justification element
    jc.set(qn('w:val'), "center")  # Sets alignment to center
    tbl_pr.append(jc) 

def fontstyle(paragraph, font_size=12, font_underline=False):
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.underline = font_underline

# Function to set paragraph line spacing
def set_paragraph_spacing(paragraph, space_after=None, space_before=None):  # 6 pt spacing after paragraph
    paragraph_format = paragraph.paragraph_format
    if space_after:
        paragraph_format.space_after = space_after
    elif space_before:
        paragraph_format.space_before=space_before


def adjust_font_size(text, max_width_inch, default_font_size=50):
    """Adjust font size based on text length to fit in one line within max_width."""
    font_size = default_font_size
    estimated_width = len(text) * (font_size * 0.002)  # Simple estimation
    while estimated_width > max_width_inch:
        font_size -= 1
        estimated_width = len(text) * (font_size * 0.002)
    return font_size

def calculate_width(text, base_width_per_char=0.1):
    """
    Calculate the width in inches based on the number of characters.
    `base_width_per_char` is the estimated width in inches that each character takes.
    """
    return max(Inches(len(text) * base_width_per_char), Inches(1))  # Ensure a minimum width




header = document.sections[0].header



header_table = header.add_table(1, 2, width=Inches(1))

set_table_center(header_table)

set_no_borders(header_table)  # Remove borders

cell_logo = header_table.cell(0, 0)
logo_paragraph = cell_logo.paragraphs[0]
run_logo = logo_paragraph.add_run()

# logo_path = 'Static/pyzon.jpg'
logo_path = 'Static/logo.jpg'
institute_name = None
institute_name = "Radiance Academy"

if not institute_name:
    logo_size = Inches(5)
else:
    logo_size = Inches(1)
if logo_path:
    run_logo.add_picture(logo_path, width=logo_size)
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# institude_name = "Radiance "
# header_table.rows[0].cells[1].width = Inches(6)


if institute_name:
    cell_institute = header_table.cell(0, 1)
    estimated_width = calculate_width(institute_name, 0.36)  # Adjust base width as needed
    cell_institute.width = estimated_width

    # Add text to the cell
    cell=header_table.cell(0, 1)
    paragraph = cell.paragraphs[0]  # Access the first paragraph in the cell
    run = paragraph.add_run(institute_name)  # Add some text
    run.font.name = 'Cambria'
    run.font.size = Pt(50)
    run.underline = True
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
else:
    cell_institute = header_table.cell(0, 1)
    cell_institute.width = Inches(0)


htable = header.add_table(1, 3, width=Inches(8))
htable.style = 'Table Grid'
# Defining cells and adding text
htab_cells = htable.rows[0].cells

# First cell with exam and date
ht0 = htab_cells[0].add_paragraph("Exam : NEET")
fontstyle(ht0)
ht0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
set_paragraph_spacing(ht0, space_after=Pt(6.5))

ht0_date = htab_cells[0].add_paragraph("Date : 12-03-2024")
fontstyle(ht0_date)
ht0_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
set_paragraph_spacing(ht0_date, space_after=Pt(7))



# Second cell with exam type
ht1 = htab_cells[1].add_paragraph('Mock Test')
fontstyle(ht1)
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
set_paragraph_spacing(ht1, space_before=Pt(6.5))

ht2 = htab_cells[2].add_paragraph("Marks: 720")
fontstyle(ht2)
ht2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
set_paragraph_spacing(ht2, space_after=Pt(6.5))  # Customizable space

# Add the time as a separate paragraph for better spacing control
ht2_time = htab_cells[2].add_paragraph("Time: 3 hour")
fontstyle(ht2_time)
ht2_time.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
set_paragraph_spacing(ht2_time, space_after=Pt(4))  # Customizable space


# Save the document
document.save('docx/header_with_logo_&_name.docx')