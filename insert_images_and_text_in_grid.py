from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document

document = Document()

# def insert_images_and_text_in_grid(document, content):
#     # Create a table with the required number of rows and 2 columns
#     rows = (len(content) + 1) // 2  # Calculate rows based on the number of content items
#     table = document.add_table(rows=rows, cols=2)

#     # Set the column width and top vertical alignment for each cell
#     for row in table.rows:
#         for cell in row.cells:
#             cell.width = Cm(5)  # Set consistent width
#             cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # Align content to the top of the cell

#             # Clear any default cell padding by setting paragraph formatting
#             for paragraph in cell.paragraphs:
#                 paragraph.paragraph_format.space_before = Pt(0)
#                 paragraph.paragraph_format.space_after = Pt(0)

#     # Insert images and text
#     for i, cell in enumerate(table._cells):
#         if i < len(content):
#             item = content[i]
#             label = f"({i + 1})"  # Label for the current cell
            
#             # Add label to the cell
#             label_paragraph = cell.add_paragraph(label)
#             label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align for label
#             label_paragraph.paragraph_format.space_after = Pt(0)  # Remove extra space after label
#             label_paragraph.paragraph_format.space_before = Pt(0)  # Remove extra space before label
            
#             if item.startswith('img:'):  # Check if the item is an image path
#                 image_path = item[4:]  # Remove 'img:' prefix
                
#                 # Add the image below the label
#                 try:
#                     image_paragraph = cell.add_paragraph()
#                     run = image_paragraph.add_run()
#                     run.add_picture(image_path, width=Cm(4.0))  # Set the width of the image
#                     image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the image

#                     # Remove extra spacing from the image paragraph
#                     image_paragraph.paragraph_format.space_after = Pt(0)
#                     image_paragraph.paragraph_format.space_before = Pt(0)
#                 except Exception as e:
#                     print(f"Error inserting image at index {i}: {e}")
#             else:
#                 # Add text below the label if it's not an image
#                 text_paragraph = cell.add_paragraph("\n"+item)
#                 text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align for text
                
#                 # Remove extra spacing from the text paragraph
#                 text_paragraph.paragraph_format.space_after = Pt(0)
#                 text_paragraph.paragraph_format.space_before = Pt(0)
#         else:
#             print(f"No content found for index {i}")

content = ['img:latex_equation2.png', 'Some text for cell 2', 'img:body.png', 'Some text for cell 4']  # Mix of images and text

# Example usage

# def insert_images_and_text_in_grid(document, content):
#     # Create a table with the required number of rows and 2 columns
#     rows = (len(content) + 1) // 2  # Calculate rows based on the number of content items
#     table = document.add_table(rows=rows, cols=2)


#     # Set the column width and top vertical alignment for each cell
#     for row in table.rows:
#         for cell in row.cells:
#             cell.width = Cm(3.0)  # Set consistent width
#             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Align content to the top of the cell

#             # Clear any default cell padding by setting paragraph formatting
#             for paragraph in cell.paragraphs:
#                 paragraph.paragraph_format.space_before = Pt(0)
#                 paragraph.paragraph_format.space_after = Pt(0)

#     # Insert images and text
#     for i, cell in enumerate(table._cells):
#         if i < len(content):
#             item = content[i]
#             label = f"({i + 1})"  # Label for the current cell
            
#             # Add label to the cell
#             # label_paragraph = cell.add_paragraph(label)
#             # label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align for label
#             # label_paragraph.paragraph_format.space_after = Pt(0)  # Remove extra space after label
#             # label_paragraph.paragraph_format.space_before = Pt(0)  # Remove extra space before label
            
#             if item.startswith('text:'):  # Check if the item is text
#                 text_content = item[5:]  # Remove 'text:' prefix
#                 # Add text below the label
#                 text_paragraph = cell.add_paragraph("\n"+text_content)
#                 text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Left-align for text

#                 # Remove extra spacing from the text paragraph
#                 text_paragraph.paragraph_format.space_after = Pt(0)
#                 text_paragraph.paragraph_format.space_before = Pt(0)
#             else:
#                 # Add the image below the label if it's an image path
#                 try:
#                     image_paragraph = cell.add_paragraph(label)
#                     run = image_paragraph.add_run()
#                     run.add_picture(item, width=Cm(3.0))  # Set the width of the image
#                     image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the image

#                     # Remove extra spacing from the image paragraph
#                     image_paragraph.paragraph_format.space_after = Pt(0)
#                     image_paragraph.paragraph_format.space_before = Pt(0)
#                 except Exception as e:
#                     print(f"Error inserting image at index {i}: {e}")
#         else:
#             print(f"No content found for index {i}")

def insert_images_and_text_in_grid(document, content):
    # Create a table with the required number of rows and 2 columns
    rows = (len(content) + 1) // 2  # Calculate rows based on the number of content items
    table = document.add_table(rows=rows, cols=2)

    # Set the column width and top vertical alignment for each cell
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5)  # Set consistent width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # Align content to the top of the cell

            # Clear any default cell padding by setting paragraph formatting
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    # Insert images and text
    for i, cell in enumerate(table._cells):
        if i < len(content):
            item = content[i]
            label = f"({i + 1})"  # Label for the current cell
            
            # Add label to the cell
            # label_paragraph = cell.add_paragraph(label)
            # label_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Left-align for label
            # label_paragraph.paragraph_format.space_after = Pt(0)  # Remove extra space after label
            # label_paragraph.paragraph_format.space_before = Pt(0)  # Remove extra space before label
            
            # # Set label to normal (not bold)
            # run = label_paragraph.runs[0]
            # run.bold = False

            if item.startswith('text:'):  # Check if the item is text
                text_content = item[5:]  # Remove 'text:' prefix
                # Add text below the label
                text_paragraph = cell.add_paragraph(label+"\n"+text_content)
                run = text_paragraph.runs[0]
                run.bold = False
                text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Left-align for text

                # Remove extra spacing from the text paragraph
                text_paragraph.paragraph_format.space_after = Pt(0)
                text_paragraph.paragraph_format.space_before = Pt(0)
            else:
                # Add the image below the label if it's an image path
                try:
                    image_paragraph = cell.add_paragraph(label)
                    run = image_paragraph.add_run()
                    run = text_paragraph.runs[0]
                    run.bold = False
                    run.add_picture(item, width=Cm(4.0))  # Set the width of the image
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the image

                    # Remove extra spacing from the image paragraph
                    image_paragraph.paragraph_format.space_after = Pt(0)
                    image_paragraph.paragraph_format.space_before = Pt(0)
                except Exception as e:
                    print(f"Error inserting image at index {i}: {e}")
        else:
            print(f"No content found for index {i}")

content = ['latex_equation2.png', 'text:Some text for cell 2', 'body.png', 'text:Some text for cell 4']  # Mix of images and text

insert_images_and_text_in_grid(document, content)
document.save('insert_images_and_text_in_grid.docx')
