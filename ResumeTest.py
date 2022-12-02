from docx import Document
from docx.enum.style import WD_STYLE_TYPE 
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.enum.section import WD_SECTION

document = Document()

# create section
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

section = document.sections[0]
# create section
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

#function for Fontstyle
def fontstyle(paragraph, font_name = 'Times New Roman', font_size = 14, font_bold = False, font_italic = False, font_underline = False):
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline

section = document.sections[0]
section = document.sections[0]
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '2')

paragraph=document.add_heading("Junaid Ahmad Chaudhary")
paragraph.alignment=0
font = paragraph.style.font
font.size=Pt(25)
font.underline = True

par=document.add_paragraph('\n+91 987654321\njc.junaid.chaudhary@gmail.com\nMumbai,Maharashtra,India')
par.alignment=0
font = paragraph.style.font
# font.size=Pt(14)

pr=document.add_picture('static/dp.png', width=Inches(2))
pr.alignment=0


# fontstyle(paragraph, font_size=28,font_underline=True)
main_section = document.add_section(0)
main_section.top_margin=Inches(0.3)
sectPr = main_section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'), '1')
#header table
htable=document.add_table(2,2)

htab_cells=htable.rows[0].cells
htab_cells=htable.rows[1].cells

ht0=htab_cells[0].add_paragraph("Junaid Ahmad Chaudhary")
ht0.add_run('bold').bold = True
ht0.add_run('italic...').italic = True
ht0.add_run('italic.').font.size = 20


# fontstyle(ht0,font_size = 25)


ht1=htab_cells[1].add_paragraph('987654321\njc.junaid.chaudhary@gmail.com\nMumbai,Maharashtra,India')
# fontstyle(ht1)
ht1.add_run('bold').bold = True
ht1.add_run('italic...').italic = True
ht1.add_run('italic.').font.size = 13
ht1.alignment = 0

# ht2=htab_cells[2].add_paragraph("987654321")
# ht2.alignment = 0

line=document.add_paragraph("_")



# .................................
q=document.add_heading('Document Title', 0) 
q.add_run("ho")


graduation = document.add_paragraph()
graduation.add_run('Bachelors of Engineering Information Technology').bold = True
graduation.add_run('(2023)').bold = True
graduation.add_run('\nM.H Saboo Siddik College of Engineering | CGPA: 8.29')

hsc = document.add_paragraph()
hsc.add_run('Ramniranjan Jhunjhunwala college | ').bold = True
hsc.add_run('HSC|(2019)').bold = True
hsc.add_run('Score:64')

ssc = document.add_paragraph()
ssc.add_run('Yogiraj Shree Krishna Vidyalaya | ').bold = True
ssc.add_run('SSC | (2017)').bold = True
ssc.add_run('core:89')


document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)




table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'

row_cells = table.add_row().cells
row_cells[0].picture = document.add_picture('static/dp.png', width=Inches(1.25))
row_cells[1].text = "junaid"
    


document.add_page_break()

document.save('docx\Resume Test1.docx')